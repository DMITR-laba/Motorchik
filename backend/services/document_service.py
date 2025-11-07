import os
import uuid
import hashlib
import mimetypes
from typing import List, Optional, Tuple, BinaryIO
from sqlalchemy.orm import Session
from sqlalchemy import and_, or_, func
from models.database import Document, Category, Tag, DocumentChunk
from models.schemas import DocumentCreate, DocumentUpdate
import requests
import time
from app.core.config import settings


class DocumentService:
    def __init__(self, db: Session):
        self.db = db
    
    def create_document(self, document_data: DocumentCreate, file_content: bytes) -> Document:
        """Создает новый документ в базе данных"""
        # Проверяем, не существует ли уже документ с таким же именем файла
        existing_document = self.db.query(Document).filter(
            Document.original_filename == document_data.original_filename
        ).first()
        
        if existing_document:
            raise ValueError(f"Документ с именем '{document_data.original_filename}' уже существует")
        
        # Генерируем уникальное имя файла
        file_hash = hashlib.md5(file_content).hexdigest()
        filename = f"{file_hash}_{uuid.uuid4().hex[:8]}.{document_data.file_type}"
        
        # Создаем документ
        document = Document(
            filename=filename,
            original_filename=document_data.original_filename,
            file_type=document_data.file_type,
            file_size=len(file_content),
            file_content=file_content,
            language=document_data.language,
            path=document_data.path,
            processing_status="pending"
        )
        
        self.db.add(document)
        self.db.flush()  # Получаем ID
        
        # Добавляем категории
        if document_data.category_ids:
            categories = self.db.query(Category).filter(Category.id.in_(document_data.category_ids)).all()
            document.categories.extend(categories)
        
        # Добавляем теги
        if document_data.tag_names:
            for tag_name in document_data.tag_names:
                tag = self.db.query(Tag).filter(Tag.name == tag_name).first()
                if not tag:
                    tag = Tag(name=tag_name)
                    self.db.add(tag)
                    self.db.flush()
                document.tags.append(tag)
        
        self.db.commit()
        self.db.refresh(document)
        return document
    
    def get_documents(self, skip: int = 0, limit: int = 100, search: Optional[str] = None) -> Tuple[List[Document], int]:
        """Получает список документов с пагинацией и поиском"""
        query = self.db.query(Document)
        
        if search:
            search_filter = or_(
                Document.original_filename.ilike(f"%{search}%"),
                Document.title.ilike(f"%{search}%"),
                Document.topic.ilike(f"%{search}%"),
                Document.extracted_text.ilike(f"%{search}%")
            )
            query = query.filter(search_filter)
        
        total = query.count()
        documents = query.offset(skip).limit(limit).all()
        return documents, total
    
    def get_document(self, document_id: int) -> Optional[Document]:
        """Получает документ по ID"""
        return self.db.query(Document).filter(Document.id == document_id).first()
    
    def update_document(self, document_id: int, document_data: DocumentUpdate) -> Optional[Document]:
        """Обновляет документ"""
        document = self.get_document(document_id)
        if not document:
            return None
        
        # Обновляем поля
        if document_data.title is not None:
            document.title = document_data.title
        if document_data.topic is not None:
            document.topic = document_data.topic
        if document_data.summary is not None:
            document.summary = document_data.summary
        if hasattr(document_data, 'path') and document_data.path is not None:
            document.path = document_data.path
        
        # Обновляем категории
        if document_data.category_ids is not None:
            document.categories.clear()
            categories = self.db.query(Category).filter(Category.id.in_(document_data.category_ids)).all()
            document.categories.extend(categories)
        
        # Обновляем теги
        if document_data.tag_names is not None:
            document.tags.clear()
            for tag_name in document_data.tag_names:
                tag = self.db.query(Tag).filter(Tag.name == tag_name).first()
                if not tag:
                    tag = Tag(name=tag_name)
                    self.db.add(tag)
                    self.db.flush()
                document.tags.append(tag)
        
        self.db.commit()
        self.db.refresh(document)
        return document
    
    def delete_document(self, document_id: int) -> bool:
        """Удаляет документ"""
        document = self.get_document(document_id)
        if not document:
            return False
        
        self.db.delete(document)
        self.db.commit()
        return True
    
    def process_document(self, document_id: int) -> bool:
        """Обрабатывает документ: извлекает текст, генерирует тему и теги"""
        document = self.get_document(document_id)
        if not document:
            return False
        
        try:
            # Обновляем статус
            document.processing_status = "processing"
            self.db.commit()
            
            # Извлекаем текст из файла
            extracted_text = self._extract_text_from_file(document.file_content, document.file_type)
            if not extracted_text:
                raise Exception("Не удалось извлечь текст из файла")
            
            document.extracted_text = extracted_text
            
            # Генерируем тему документа
            topic = self._generate_topic(extracted_text)
            document.topic = topic
            
            # Генерируем заголовок
            title = self._generate_title(extracted_text, document.original_filename)
            document.title = title
            
            # Генерируем категории
            categories = self._generate_categories(extracted_text, title)
            if categories:
                # Очищаем старые категории и добавляем новые
                document.categories.clear()
                for category_name in categories:
                    category = self.db.query(Category).filter(Category.name == category_name).first()
                    if not category:
                        category = Category(name=category_name)
                        self.db.add(category)
                        self.db.flush()
                    document.categories.append(category)
            
            # Генерируем теги
            tags = self._generate_tags(extracted_text, title)
            if tags:
                # Очищаем старые теги и добавляем новые
                document.tags.clear()
                for tag_name in tags:
                    tag = self.db.query(Tag).filter(Tag.name == tag_name).first()
                    if not tag:
                        tag = Tag(name=tag_name)
                        self.db.add(tag)
                        self.db.flush()
                    document.tags.append(tag)
            
            # Генерируем краткое содержание
            summary = self._generate_summary(extracted_text)
            document.summary = summary
            
            # Создаем чанки текста с эмбеддингами
            chunks = self.create_document_chunks(document.id)
            
            # Обновляем статус и время обработки
            document.processing_status = "completed"
            document.processed_at = func.now()
            document.error_message = None
            
            self.db.commit()
            return True
            
        except Exception as e:
            document.processing_status = "failed"
            document.error_message = str(e)
            self.db.commit()
            return False
    
    def _extract_text_from_file(self, file_content: bytes, file_type: str) -> str:
        """Извлекает текст из файла в зависимости от типа"""
        try:
            if file_type.lower() == 'txt':
                return file_content.decode('utf-8', errors='ignore')
            
            elif file_type.lower() == 'pdf':
                return self._extract_text_from_pdf(file_content)
            
            elif file_type.lower() in ['doc', 'docx']:
                return self._extract_text_from_doc(file_content)
            
            else:
                raise Exception(f"Неподдерживаемый тип файла: {file_type}")
                
        except Exception as e:
            raise Exception(f"Ошибка извлечения текста: {str(e)}")
    
    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Извлекает текст из PDF файла"""
        try:
            import PyPDF2
            import io
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except ImportError:
            # Fallback: используем внешний сервис или возвращаем ошибку
            raise Exception("PyPDF2 не установлен. Установите: pip install PyPDF2")
        except Exception as e:
            raise Exception(f"Ошибка чтения PDF: {str(e)}")
    
    def _extract_text_from_doc(self, file_content: bytes) -> str:
        """Извлекает текст из DOC/DOCX файла"""
        try:
            import docx
            import io
            
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except ImportError:
            raise Exception("python-docx не установлен. Установите: pip install python-docx")
        except Exception as e:
            raise Exception(f"Ошибка чтения DOC/DOCX: {str(e)}")
    
    def _generate_topic(self, text: str) -> str:
        """Генерирует тему документа с помощью Mistral AI"""
        try:
            prompt = f"""
            Проанализируй следующий текст и определи основную тему документа в 2-3 словах.
            Отвечай только темой, без дополнительных объяснений.
            
            Текст: {text[:2000]}
            """
            
            url = f"{settings.mistral_base_url}/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {settings.mistral_api_key}",
                "Content-Type": "application/json",
            }
            payload = {
                "model": settings.mistral_model,
                "messages": [
                    {"role": "system", "content": "Ты помощник для анализа документов. Отвечай кратко и по существу."},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.3,
                "max_tokens": 50,
                "stream": False,
            }
            
            for attempt in range(3):
                try:
                    resp = requests.post(url, headers=headers, json=payload, timeout=60)
                    if resp.status_code == 429:
                        time.sleep(1 * (2 ** attempt))
                        continue
                    resp.raise_for_status()
                    data = resp.json() or {}
                    choices = data.get("choices") or []
                    if choices:
                        topic = (choices[0].get("message") or {}).get("content", "").strip()
                        return topic[:100] if topic else "Неизвестная тема"
                    return "Неизвестная тема"
                except Exception:
                    if attempt == 2:
                        return "Неизвестная тема"
                    time.sleep(1)
            
        except Exception:
            return "Неизвестная тема"
    
    def _generate_title(self, text: str, original_filename: str) -> str:
        """Генерирует заголовок документа"""
        try:
            # Используем первые 200 символов текста или имя файла
            if len(text) > 200:
                return text[:200] + "..."
            elif text.strip():
                return text.strip()
            else:
                return original_filename
        except:
            return original_filename
    
    def _generate_summary(self, text: str) -> str:
        """Генерирует краткое содержание документа"""
        try:
            prompt = f"""
            Создай краткое содержание следующего документа в 2-3 предложениях.
            Выдели основные моменты и ключевые идеи.
            
            Текст: {text[:2000]}
            """
            
            url = f"{settings.mistral_base_url}/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {settings.mistral_api_key}",
                "Content-Type": "application/json",
            }
            payload = {
                "model": settings.mistral_model,
                "messages": [
                    {"role": "system", "content": "Ты помощник для создания кратких содержаний. Отвечай на русском языке."},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.3,
                "max_tokens": 200,
                "stream": False,
            }
            
            for attempt in range(3):
                try:
                    resp = requests.post(url, headers=headers, json=payload, timeout=60)
                    if resp.status_code == 429:
                        time.sleep(1 * (2 ** attempt))
                        continue
                    resp.raise_for_status()
                    data = resp.json() or {}
                    choices = data.get("choices") or []
                    if choices:
                        summary = (choices[0].get("message") or {}).get("content", "").strip()
                        return summary[:500] if summary else ""
                    return ""
                except Exception:
                    if attempt == 2:
                        return ""
                    time.sleep(1)
            
        except Exception:
            return ""
    
    def _generate_categories(self, text: str, title: str = "") -> List[str]:
        """Генерирует категории для документа с помощью Mistral AI"""
        try:
            prompt = f"""
            Проанализируй следующий документ и определи 1-3 наиболее подходящие категории из списка:
            - Техническая документация
            - Бизнес-процессы
            - Правовые документы
            - Научные статьи
            - Инструкции
            - Отчеты
            - Презентации
            - Финансовые документы
            - Маркетинг
            - Образование
            - Здоровье
            - Технологии
            - Управление
            - Другое
            
            Заголовок: {title}
            Текст: {text[:1500]}
            
            Верни только названия категорий через запятую, без дополнительных объяснений.
            """
            
            url = f"{settings.mistral_base_url}/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {settings.mistral_api_key}",
                "Content-Type": "application/json",
            }
            payload = {
                "model": settings.mistral_model,
                "messages": [
                    {"role": "system", "content": "Ты помощник для категоризации документов. Отвечай на русском языке."},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.3,
                "max_tokens": 100,
                "stream": False,
            }
            
            for attempt in range(3):
                try:
                    resp = requests.post(url, headers=headers, json=payload, timeout=60)
                    if resp.status_code == 429:
                        time.sleep(1 * (2 ** attempt))
                        continue
                    resp.raise_for_status()
                    data = resp.json() or {}
                    choices = data.get("choices") or []
                    if choices:
                        categories_text = (choices[0].get("message") or {}).get("content", "").strip()
                        if categories_text:
                            # Разделяем по запятым и очищаем
                            categories = [cat.strip() for cat in categories_text.split(',') if cat.strip()]
                            return categories[:3]  # Максимум 3 категории
                    return []
                except Exception:
                    if attempt == 2:
                        return []
                    time.sleep(1)
            
        except Exception:
            return []
    
    def _generate_tags(self, text: str, title: str = "") -> List[str]:
        """Генерирует теги для документа с помощью Mistral AI"""
        try:
            prompt = f"""
            Проанализируй следующий документ и создай 3-7 релевантных тегов (ключевых слов).
            Теги должны быть короткими (1-3 слова), на русском языке, и отражать основное содержание.
            
            Заголовок: {title}
            Текст: {text[:1500]}
            
            Верни только теги через запятую, без дополнительных объяснений.
            """
            
            url = f"{settings.mistral_base_url}/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {settings.mistral_api_key}",
                "Content-Type": "application/json",
            }
            payload = {
                "model": settings.mistral_model,
                "messages": [
                    {"role": "system", "content": "Ты помощник для создания тегов документов. Отвечай на русском языке."},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.4,
                "max_tokens": 150,
                "stream": False,
            }
            
            for attempt in range(3):
                try:
                    resp = requests.post(url, headers=headers, json=payload, timeout=60)
                    if resp.status_code == 429:
                        time.sleep(1 * (2 ** attempt))
                        continue
                    resp.raise_for_status()
                    data = resp.json() or {}
                    choices = data.get("choices") or []
                    if choices:
                        tags_text = (choices[0].get("message") or {}).get("content", "").strip()
                        if tags_text:
                            # Разделяем по запятым и очищаем
                            tags = [tag.strip() for tag in tags_text.split(',') if tag.strip()]
                            return tags[:7]  # Максимум 7 тегов
                    return []
                except Exception:
                    if attempt == 2:
                        return []
                    time.sleep(1)
            
        except Exception:
            return []
    
    def search_documents_for_rag(self, query: str, limit: int = 5) -> List[Document]:
        """Поиск документов для RAG системы"""
        search_filter = or_(
            Document.title.ilike(f"%{query}%"),
            Document.topic.ilike(f"%{query}%"),
            Document.extracted_text.ilike(f"%{query}%"),
            Document.summary.ilike(f"%{query}%")
        )
        return self.db.query(Document).filter(
            and_(search_filter, Document.processing_status == "completed")
        ).limit(limit).all()
    
    def _split_text_into_chunks(self, text: str, chunk_size: int = 500) -> List[str]:
        """Разбивает текст на чанки заданного размера"""
        if not text:
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Если это не последний чанк, ищем ближайший пробел или перенос строки
            if end < len(text):
                # Ищем последний пробел или перенос строки в пределах чанка
                for i in range(end, start + chunk_size // 2, -1):
                    if text[i] in [' ', '\n', '\t', '.', '!', '?']:
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end
        
        return chunks
    
    def _build_embedding_context(self, text: str, context: dict = None) -> str:
        """Строит контекстный текст для эмбеддинга"""
        if not context:
            return text
        
        context_parts = []
        
        # Добавляем название файла
        if context.get('filename'):
            context_parts.append(f"Файл: {context['filename']}")
        
        # Добавляем путь
        if context.get('path'):
            context_parts.append(f"Путь: {context['path']}")
        
        # Добавляем тему документа
        if context.get('topic'):
            context_parts.append(f"Тема: {context['topic']}")
        
        # Добавляем категории
        if context.get('categories'):
            categories_str = ", ".join(context['categories'])
            context_parts.append(f"Категории: {categories_str}")
        
        # Добавляем теги
        if context.get('tags'):
            tags_str = ", ".join(context['tags'])
            context_parts.append(f"Теги: {tags_str}")
        
        # Добавляем сам текст чанка
        context_parts.append(f"Содержание: {text}")
        
        return " | ".join(context_parts)
    
    def _generate_embedding(self, text: str, context: dict = None) -> Optional[str]:
        """Генерирует эмбеддинг для текста с контекстной информацией"""
        try:
            # Формируем контекстный текст для эмбеддинга
            context_text = self._build_embedding_context(text, context)
            
            url = f"{settings.mistral_base_url}/v1/embeddings"
            headers = {
                "Authorization": f"Bearer {settings.mistral_api_key}",
                "Content-Type": "application/json",
            }
            payload = {
                "model": "mistral-embed",  # Используем модель для эмбеддингов
                "input": context_text[:1000],  # Ограничиваем размер текста
            }
            
            for attempt in range(3):
                try:
                    resp = requests.post(url, headers=headers, json=payload, timeout=30)
                    if resp.status_code == 429:
                        time.sleep(1 * (2 ** attempt))
                        continue
                    resp.raise_for_status()
                    data = resp.json()
                    if "data" in data and len(data["data"]) > 0:
                        embedding = data["data"][0]["embedding"]
                        return str(embedding)  # Сохраняем как JSON строку
                    return None
                except Exception:
                    if attempt == 2:
                        return None
                    time.sleep(1)
            
        except Exception:
            return None
    
    def create_document_chunks(self, document_id: int, show_progress: bool = False) -> List[DocumentChunk]:
        """Создает чанки для документа"""
        document = self.get_document(document_id)
        if not document or not document.extracted_text:
            return []
        
        # Удаляем существующие чанки
        self.db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
        
        # Разбиваем текст на чанки
        chunks = self._split_text_into_chunks(document.extracted_text)
        created_chunks = []
        
        # Собираем контекстную информацию для эмбеддингов
        context = {
            'filename': document.original_filename,
            'path': document.path,
            'topic': document.topic,
            'categories': [cat.name for cat in document.categories],
            'tags': [tag.name for tag in document.tags]
        }
        
        total_chunks = len(chunks)
        if show_progress:
            print(f"      📊 Создание {total_chunks} чанков с эмбеддингами...")
        
        for i, chunk_text in enumerate(chunks):
            if show_progress:
                # Показываем прогресс для эмбеддингов
                percent = (i + 1) / total_chunks * 100
                bar_length = 30
                filled = int(bar_length * (i + 1) // total_chunks)
                bar = '█' * filled + '░' * (bar_length - filled)
                print(f'\r      🔄 Эмбеддинги |{bar}| {percent:.1f}% ({i+1}/{total_chunks})', end='', flush=True)
            
            # Генерируем эмбеддинг для чанка с контекстом
            embedding = self._generate_embedding(chunk_text, context)
            
            # Создаем чанк
            chunk = DocumentChunk(
                document_id=document_id,
                chunk_index=i,
                text=chunk_text,
                embedding=embedding
            )
            
            self.db.add(chunk)
            created_chunks.append(chunk)
        
        if show_progress:
            print()  # Новая строка после прогресс-бара
        
        self.db.commit()
        
        # Обновляем статус документа
        document.processing_status = "completed"
        self.db.commit()
        
        return created_chunks
    
    def get_document_chunks(self, document_id: int) -> List[DocumentChunk]:
        """Получает чанки документа"""
        return self.db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id
        ).order_by(DocumentChunk.chunk_index).all()
    
    def search_document_chunks(self, query: str, limit: int = 10) -> List[DocumentChunk]:
        """Поиск по чанкам документов"""
        # Простой текстовый поиск по чанкам
        search_filter = DocumentChunk.text.ilike(f"%{query}%")
        return self.db.query(DocumentChunk).filter(search_filter).limit(limit).all()
